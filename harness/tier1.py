#!/usr/bin/env python3
# SPDX-License-Identifier: MIT OR Apache-2.0
"""Tier 1 transport-survivability driver: pump the carrier vectors emitted by the
Rust benchmark through *real* text pipelines, then hand the results back to the
Rust classifier.

Unlike the Tier 0 probes (deterministic filters we wrote), these transports are
production libraries and system tools applied as-is:

  iconv-ascii-translit   UTF-8 -> ASCII//TRANSLIT (drops non-ASCII)
  pandoc-md-html-md      Markdown -> HTML -> Markdown round trip
  textutil-rtf           plain text -> RTF -> plain text (macOS Rich Text engine)
  bleach-sanitize        Mozilla bleach HTML sanitizer
  nh3-sanitize           nh3 (ammonia) HTML sanitizer
  ftfy-fix-text          ftfy text repair

A transport whose tool or library is missing is skipped and reported, never
silently treated as a pass.

Run with the harness venv so bleach/nh3/ftfy import:
  harness/.venv/bin/python harness/tier1.py
"""

import json
import os
import shutil
import subprocess
import sys
import tempfile

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
EXAMPLE = ["cargo", "run", "-q", "--release", "--example", "transport_survivability"]


def emit_vectors():
    out = subprocess.run(
        EXAMPLE + ["emit"], cwd=ROOT, capture_output=True, text=True, check=True
    )
    return json.loads(out.stdout)


def classify(results):
    with tempfile.NamedTemporaryFile("w", suffix=".json", delete=False) as f:
        json.dump(results, f)
        path = f.name
    try:
        subprocess.run(EXAMPLE + ["classify", path], cwd=ROOT, check=True)
    finally:
        os.unlink(path)


def iconv_translit(s):
    p = subprocess.run(
        ["iconv", "-f", "UTF-8", "-t", "ASCII//TRANSLIT"],
        input=s,
        capture_output=True,
        text=True,
    )
    return p.stdout


def pandoc_roundtrip(s):
    html = subprocess.run(
        ["pandoc", "-f", "markdown", "-t", "html"],
        input=s,
        capture_output=True,
        text=True,
        check=True,
    ).stdout
    return subprocess.run(
        ["pandoc", "-f", "html", "-t", "markdown"],
        input=html,
        capture_output=True,
        text=True,
        check=True,
    ).stdout


def textutil_rtf(s):
    with tempfile.TemporaryDirectory() as d:
        a, r, b = (os.path.join(d, n) for n in ("a.txt", "a.rtf", "b.txt"))
        with open(a, "w") as f:
            f.write(s)
        subprocess.run(["textutil", "-convert", "rtf", a, "-output", r], check=True)
        subprocess.run(["textutil", "-convert", "txt", r, "-output", b], check=True)
        with open(b) as f:
            return f.read()


def json_roundtrip(s):
    return json.loads(json.dumps(s))


def sqlite_roundtrip(s):
    import sqlite3

    con = sqlite3.connect(":memory:")
    con.execute("create table t(x text)")
    con.execute("insert into t values (?)", (s,))
    (out,) = con.execute("select x from t").fetchone()
    con.close()
    return out


def email_roundtrip(s):
    import email
    import email.policy
    from email.message import EmailMessage

    m = EmailMessage()
    m.set_content(s)
    m2 = email.message_from_bytes(m.as_bytes(), policy=email.policy.default)
    return m2.get_content()


def markdown_pipeline(s):
    import html2text
    import markdown

    return html2text.html2text(markdown.markdown(s))


def lxml_html_roundtrip(s):
    import lxml.html

    return lxml.html.fragment_fromstring(f"<p>{s}</p>").text_content()


def docx_roundtrip(s):
    import docx

    with tempfile.TemporaryDirectory() as d:
        path = os.path.join(d, "a.docx")
        doc = docx.Document()
        doc.add_paragraph(s)
        doc.save(path)
        return "\n".join(p.text for p in docx.Document(path).paragraphs)


def tidy_html(s):
    import html2text

    p = subprocess.run(
        ["tidy", "-q", "--show-warnings", "no", "--force-output", "yes"],
        input=f"<p>{s}</p>",
        capture_output=True,
        text=True,
    )
    return html2text.html2text(p.stdout)


def build_transports():
    transports, skipped = [], []

    def add_tool(name, tool, fn):
        if shutil.which(tool):
            transports.append((name, fn))
        else:
            skipped.append(name)

    add_tool("iconv-ascii-translit", "iconv", iconv_translit)
    add_tool("pandoc-md-html-md", "pandoc", pandoc_roundtrip)
    add_tool("textutil-rtf", "textutil", textutil_rtf)

    try:
        import bleach

        transports.append(
            ("bleach-sanitize", lambda s: bleach.clean(s, tags=[], strip=True))
        )
    except ImportError:
        skipped.append("bleach-sanitize")
    try:
        import nh3

        transports.append(("nh3-sanitize", lambda s: nh3.clean(s, tags=set())))
    except ImportError:
        skipped.append("nh3-sanitize")
    try:
        import ftfy

        transports.append(("ftfy-fix-text", ftfy.fix_text))
    except ImportError:
        skipped.append("ftfy-fix-text")

    transports.append(("json-roundtrip", json_roundtrip))
    transports.append(("sqlite-roundtrip", sqlite_roundtrip))
    transports.append(("email-mime-roundtrip", email_roundtrip))

    for name, fn, mod in [
        ("markdown-pipeline", markdown_pipeline, "html2text"),
        ("lxml-html-parse", lxml_html_roundtrip, "lxml.html"),
        ("docx-roundtrip", docx_roundtrip, "docx"),
    ]:
        try:
            __import__(mod)
            transports.append((name, fn))
        except ImportError:
            skipped.append(name)

    if shutil.which("tidy"):
        try:
            __import__("html2text")
            transports.append(("tidy-html", tidy_html))
        except ImportError:
            skipped.append("tidy-html")
    else:
        skipped.append("tidy-html")

    return transports, skipped


def main():
    vectors = emit_vectors()
    transports, skipped = build_transports()

    results = {}
    for tname, tf in transports:
        per = {}
        for method, embedded in vectors.items():
            try:
                per[method] = tf(embedded)
            except Exception as e:  # noqa: BLE001 - report, do not fake a pass
                print(f"  {tname}/{method}: transport error: {e}", file=sys.stderr)
        results[tname] = per

    print(f"real transports run: {[t for t, _ in transports]}", file=sys.stderr)
    if skipped:
        print(f"skipped (tool/lib missing): {skipped}", file=sys.stderr)
    print("\n== Tier 1: real transports ==", file=sys.stderr)
    classify(results)


if __name__ == "__main__":
    main()
